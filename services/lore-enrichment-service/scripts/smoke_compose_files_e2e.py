"""Live E2E for Compose Slice 3 (mode F — attach files + OCR).

Generates three files in-process — a .txt + a .docx (text-layer) + an image-only
"scanned" PDF (Pillow image, English text, OCR'd by Tesseract) — uploads each,
polls extraction to ready, asserts text/OCR, then POST /compose input_source=files
→ grounded quarantined proposal.

Usage:
  python scripts/smoke_compose_files_e2e.py --base http://127.0.0.1:8221 \
    --project <book> --user <uuid> --jwt-secret <secret> --embed <ref> --gen <ref>
"""

from __future__ import annotations

import argparse
import io
import time

import docx
import httpx
import jwt as pyjwt
from PIL import Image, ImageDraw, ImageFont

TARGET = "蓬萊仙島-FILES-SMOKE"


def _txt() -> bytes:
    return ("蓬萊乃東海之上的仙山，雲霧繚繞，仙人所居，服其瓊草可長生。" * 3).encode("utf-8")


def _docx() -> bytes:
    d = docx.Document()
    d.add_paragraph("第一段：蓬萊宮闕，金玉砌成。")
    d.add_paragraph("第二段：仙人於此煉丹修道。")
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _scanned_pdf() -> bytes:
    # image-only PDF (no text layer) → forces the OCR path. English text avoids a
    # CJK-font dependency on the host; the service OCRs with chi_sim+chi_tra+eng.
    img = Image.new("RGB", (1400, 400), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.load_default(size=52)
    except TypeError:  # Pillow < 10 load_default takes no size
        font = ImageFont.load_default()
    draw.text((40, 160), "Penglai is a sacred island in the eastern sea.", fill="black", font=font)
    b = io.BytesIO()
    img.save(b, format="PDF")
    return b.getvalue()


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--base", required=True)
    ap.add_argument("--project", required=True)
    ap.add_argument("--user", required=True)
    ap.add_argument("--jwt-secret", required=True)
    ap.add_argument("--embed", required=True)
    ap.add_argument("--gen", required=True)
    args = ap.parse_args()
    base = args.base.rstrip("/")
    uh = {"Authorization": f"Bearer {pyjwt.encode({'sub': args.user}, args.jwt_secret, algorithm='HS256')}"}

    files = [
        ("ref.txt", _txt(), "text/plain", False),
        ("ref.docx", _docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", False),
        ("scan.pdf", _scanned_pdf(), "application/pdf", True),  # expect OCR
    ]

    with httpx.Client(timeout=60.0) as c:
        upload_ids: list[str] = []
        for name, data, mime, expect_ocr in files:
            print(f"=== upload {name} ===")
            r = c.post(
                f"{base}/v1/lore-enrichment/uploads", headers=uh,
                files={"file": (name, data, mime)},
                data={"book_id": args.project, "project_id": args.project, "license_asserted": "public_domain"},
            )
            print(r.status_code, r.text[:160])
            if r.status_code != 202:
                print(f"FAIL: upload {name} not accepted"); return 1
            uid = r.json()["upload_id"]
            # poll to ready
            status, view = "processing", {}
            for _ in range(40):
                pr = c.get(f"{base}/v1/lore-enrichment/uploads/{uid}", headers=uh)
                view = pr.json()
                status = view.get("status")
                if status in ("ready", "failed"):
                    break
                time.sleep(1.5)
            print(f"  {name}: status={status} chars={view.get('extracted_chars')} ocr={view.get('ocr_used')}")
            if status != "ready":
                print(f"FAIL: {name} did not extract (status={status}, err={view.get('error')})"); return 1
            if (view.get("extracted_chars") or 0) <= 0:
                print(f"FAIL: {name} extracted no text"); return 1
            if expect_ocr and not view.get("ocr_used"):
                print(f"FAIL: {name} expected OCR but ocr_used=False (Tesseract/poppler missing?)"); return 1
            upload_ids.append(uid)
        print(f"all uploads ready: {len(upload_ids)} (incl. OCR'd scan) ✓")

        print("\n=== POST /compose input_source=files ===")
        r = c.post(
            f"{base}/v1/lore-enrichment/projects/{args.project}/compose", headers=uh,
            json={
                "book_id": args.project, "input_source": "files", "upload_ids": upload_ids,
                "target": {"mode": "new", "canonical_name": TARGET, "entity_kind": "location"},
                "embedding_model_ref": args.embed, "generation_model_ref": args.gen,
                "technique": "retrieval", "max_spend_tokens": 100000, "top_k": 5,
            },
        )
        print(r.status_code, r.text[:200])
        if r.status_code != 202:
            print("FAIL: compose files did not enqueue"); return 1
        job_id = r.json()["job_id"]

        print("\n=== poll job → completed ===")
        status_val = None
        for _ in range(60):
            jr = c.get(f"{base}/v1/lore-enrichment/jobs?book_id={args.project}&limit=50", headers=uh)
            job = next((j for j in jr.json().get("items", []) if j.get("job_id") == job_id), None)
            status_val = (job or {}).get("status")
            print(f"  job status={status_val}")
            if status_val in ("completed", "failed", "error", "done"):
                break
            time.sleep(3)
        if status_val not in ("completed", "done"):
            print(f"FAIL/INCONCLUSIVE: job ended status={status_val}"); return 2

        pr = c.get(f"{base}/v1/lore-enrichment/proposals?book_id={args.project}&limit=50", headers=uh)
        mine = [p for p in pr.json().get("items", []) if TARGET in str(p.get("canonical_name") or p.get("title") or "")]
        if not mine:
            print("FAIL: no proposal for the files target"); return 1
        p = mine[0]
        assert (p.get("review_status") or p.get("status")) in ("proposed", "pending"), f"not quarantined: {p}"
        conf = p.get("confidence")
        assert conf is None or conf < 1.0, f"canon confidence: {conf}"
        print(f"  proposal {p.get('proposal_id') or p.get('id')} quarantined (conf={conf}) ✓")
        # cleanup: reject the quarantined proposal (glossary was never touched — new target)
        c.post(f"{base}/v1/lore-enrichment/proposals/{p.get('proposal_id') or p.get('id')}/reject"
               f"?project_id={args.project}", headers=uh, json={"reason": "S3 files smoke cleanup"})

    print("\nSMOKE OK — upload(.txt+.docx+scanned.pdf w/ OCR) → extract → ingest → "
          "grounded quarantined proposal; cleaned up.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
